from lib2to3.pgen2.driver import Driver
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
import pyperclip
from docx import Document
from selenium.webdriver.common.keys import Keys
from posixpath import expanduser
from selenium.webdriver.support.ui import Select
from PIL import Image
from io import BytesIO
from selenium.webdriver.common.action_chains import ActionChains
from datetime import date, timedelta
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC


BEG = date.today() - timedelta(days=61)
END = date.today() - timedelta(days=31)
d1 = BEG.strftime("%b %d, %Y")
d2 = END.strftime("%b %d, %Y")
print(d1)
print(d2)


SHIPMENTID = pyperclip.paste()
pyperclip.copy(SHIPMENTID)

PATH = "C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\\webdriver\\chromedriver.exe"
driver = webdriver.Chrome(PATH)

options = webdriver.ChromeOptions() 
options.add_argument("user-data-dir=C:\\Users\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\webdriver\\newchrome")  #Path to your chrome profile

driver = webdriver.Chrome(executable_path="C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\\webdriver\\chromedriver.exe", chrome_options=options)

time.sleep(1)

WEB = driver.get("https://scm.controlant.com/global/shipments/search/" + SHIPMENTID + "/chart?page=1")
driver.maximize_window()


#wait for page to load and element appear
try:
    element = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, "//label[normalize-space()='Light events']")))
except:
    driver.quit()
#time.sleep(10)

driver.find_element(By.XPATH, "//label[normalize-space()='Light events']").click()

time.sleep(5)

driver.find_element(By.XPATH, "//button[@class='multi-select-button ng-binding']").click()
LOGGERID1 = driver.find_element(By.XPATH, "//span[@class='item-label ng-binding']").text

image_chart = driver.find_element(By.XPATH, "//div[@class='pane-frame']").screenshot("C:\\Users\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\test.png")

driver.find_element(By.XPATH, "//a[normalize-space()='Info']").click()

time.sleep(1)

#data exctract
SHIPPED = driver.find_element(By.CSS_SELECTOR, ".ng-binding.ng-scope[data-testid='span-shipment-shipped']").text

ch = "2022"    
stripped1 = SHIPPED.strip("Shipped")
head, sep, tail = stripped1.partition(ch)
DATEBEG = head + sep
print(DATEBEG)
DELIVERED = driver.find_element(By.CSS_SELECTOR, ".ng-binding.ng-scope[data-testid='span-shipment-delivered']").text

ch = "2022"
stripped2 = DELIVERED.strip("Delivered")
head, sep, tail = stripped2.partition(ch)
DATEEND = head + sep
print(DATEEND)

OBD = driver.find_element(By.CSS_SELECTOR, "h2[class='ng-binding']").text
print(OBD)
ORIGIN = driver.find_element(By.CSS_SELECTOR, "div[class='origin'] strong[class='ng-binding ng-scope']").text
DESTINATION = driver.find_element(By.CSS_SELECTOR, "div[class='destination'] strong[class='ng-binding ng-scope']").text

driver.find_element(By.CSS_SELECTOR, ".logger-link.ng-binding.ng-scope").click()
time.sleep(1)
driver.find_element(By.XPATH, "//a[normalize-space()='Info']").click()
time.sleep(2)
LASTCHIN = driver.find_element(By.CSS_SELECTOR, "strong[class='ng-binding ng-scope']").text
LOGGERID = driver.find_element(By.CSS_SELECTOR, "h2[class='ng-binding'] span[class='logger-id ng-binding']").text







# >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>> second part


#PATH = "C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\\webdriver\\chromedriver.exe"
#driver = webdriver.Chrome(PATH)

#options = webdriver.ChromeOptions() 
#options.add_argument("user-data-dir=C:\\Users\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\webdriver\\newchrome2")  #Path to your chrome profile

#driver = webdriver.Chrome(executable_path="C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\\webdriver\\chromedriver.exe", chrome_options=options)

driver.get("https://reseller.controlant.com/search/" + LOGGERID + "/device-info")
driver.maximize_window()

#wait for page to load and element appear
try:
    element = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, By.XPATH, "/html[1]/body[1]/div[1]/main[1]/div[1]/section[1]/div[1]/div[1]/table[1]/tbody[1]/tr[11]/td[1]/input[1]")))
except:
    driver.quit()

#time.sleep(5)

#set the date
def SetTheDate():

    time.sleep(2)
    driver.find_element(By.XPATH, "/html[1]/body[1]/div[1]/main[1]/div[1]/section[1]/div[1]/div[1]/table[1]/tbody[1]/tr[11]/td[1]/input[1]").click()
    driver.find_element(By.XPATH, "//div[@view='chart data']//select[@class='ng-pristine ng-untouched ng-valid ng-not-empty']").click()

    select = Select(driver.find_element(By.XPATH, "//div[@view='chart data']//select[@class='ng-pristine ng-untouched ng-valid ng-not-empty']"))
    select.select_by_index(2)

    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").click()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").clear()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").send_keys(DATEBEG)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-end']").click()
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-end']").clear()
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-end']").send_keys(DATEEND)


#change date -/+ day
def ChangeDay():
    time.sleep(1)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-start']").send_keys(Keys.ARROW_LEFT)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//label[@for='datepicker-end'][normalize-space()='To:']").send_keys(Keys.ARROW_RIGHT + Keys.ARROW_RIGHT)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//button[@class='small'][normalize-space()='Apply']").click()


#take screenshot and crop
def Screen():
    time.sleep(5)
    LTScreen = driver.get_screenshot_as_png()
    LTCropped = Image.open(BytesIO(LTScreen))
    #LTCropped = LTCropped.crop((20, 85, 2360, 940))    #everything, but too big
    LTCropped = LTCropped.crop((20, 85, 2360, 740))     #LT only
    path2 = expanduser("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\")
    LTCropped.save(path2 + "LT Screen.png")


#LT screens for last 2 months
def LT1ScreenMonths1():
    select = Select(driver.find_element(By.XPATH, "//select[@class='ng-valid ng-not-empty ng-dirty ng-valid-parse ng-touched']"))
    select.select_by_index(1)
    time.sleep(1)
    driver.find_element(By.XPATH, "//h1[1]").click()
    LT1Screen = driver.get_screenshot_as_png()
    LT1Cropped = Image.open(BytesIO(LT1Screen))
    LT1Cropped = LT1Cropped.crop((20, 85, 2360, 740))     #LT only
    path3 = expanduser("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\")
    LT1Cropped.save(path3 + "LT1.png")

def LT2ScreenMonths2():
    time.sleep(5)
    select = Select(driver.find_element(By.XPATH, "//select[@class='ng-valid ng-not-empty ng-dirty ng-valid-parse ng-touched']"))
    select.select_by_index(2)

    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").click()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").clear()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").send_keys(d1)

    time.sleep(2)

    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//input[@id='datepicker-end']").click()
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//input[@id='datepicker-end']").clear()
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//input[@id='datepicker-end']").send_keys(d2)
    time.sleep(2)
    driver.find_element(By.XPATH, "//h1[1]").click()
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//button[@class='small'][normalize-space()='Apply']").click()
    #N = 30  # number of times you want to press TAB
    #actions = ActionChains(driver) 
    #for _ in range(N):
    #    actions = actions.send_keys(Keys.ARROW_LEFT)
    #actions.perform()

    time.sleep(2)

    LT2Screen = driver.get_screenshot_as_png()
    LT2Cropped = Image.open(BytesIO(LT2Screen))
    LT2Cropped = LT2Cropped.crop((20, 85, 2360, 740))     #LT only
    path4 = expanduser("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\")
    LT2Cropped.save(path4 + "LT2.png")

#generate file
def GenerateDocxFile():
    document = Document()

    document.add_paragraph(LASTCHIN)
    document.add_paragraph(OBD)
    document.add_paragraph(ORIGIN + " / " + DESTINATION)
    
    #document.add_paragraph(SHIPPED)
    #document.add_paragraph(DELIVERED)
    document.add_paragraph(" ")
    document.add_paragraph("Data missing:")
    document.add_paragraph(" ")

    #document.add_picture("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\scm.png")

    document.add_paragraph("Logger temperature within normal boundaries. / Logger temperature dropped below operational boundaries on ")
    
    #document.add_picture("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT Screen.png")

    document.add_paragraph("No probe failure events displayed.")

    document.save("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\screensdemo.docx")

lines = [LASTCHIN, OBD, ORIGIN + " / " + DESTINATION, " ",

 "Data missing: ", "Logger temperature within normal boundaries. / Logger temperature breached lower boundary of (-20,0°C) on ", "No probe failure events displayed."]
with open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\summary.txt", "w") as f:
    for line in lines:
        f.write(line)
        f.write('\n')

SetTheDate()
ChangeDay()
Screen()

LT1ScreenMonths1()
LT2ScreenMonths2()


with open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\summary.txt") as f:
    data = f.read()
    pyperclip.copy(data)


img1 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\test.png")
img2 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT Screen.png")
img3 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT1.png")
img4 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT2.png")
img1.show()
img2.show()
img3.show()
img4.show()

